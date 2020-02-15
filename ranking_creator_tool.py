from docx.api import Document
from docx.shared import Pt
from docx.enum.table import WD_ROW_HEIGHT
from pprint import pprint 
from collections import OrderedDict 
from datetime import datetime


# Reading .docx table and parsing it into pythons ordered dict.

document = Document('12.01.2020 KLASYFIKACJA IV GPSWA MĘŻCZYZN.docx')
table = document.tables[0]
data = {}
keys = []
ints_lst = [0, 3, 4]
for i, row in enumerate(table.rows):
    text = [cell.text for cell in row.cells]
    if i == 0:
        keys = tuple(text)
        continue
    temp_key = ''
    to_parse = []
    for j, cell in enumerate(row.cells):
        if j in ints_lst:
            to_parse.append(int(cell.text))
        elif j == 2:
            temp_key = cell.text
            to_parse.append(cell.text)
        else:
            to_parse.append(cell.text)
    row_data = dict(zip(keys, to_parse))
    data[temp_key] = row_data

# '-',  '↓ 1',  '↑ 3',

# Functions
def add_new_player(pd, dd):
    # adds new player to ranking table
    # allocates points and number of tournaments played
    # pd - player_data , dd - data_dictionary
    dd[pd[1]] = {
        'LRT': 1,
        'PUNKTY': tounament_type[pd[0]],
        'RANKING': 0,
        'RUCH': '-',
        'ZAWODNIK': pd[1]}

def update_existing_player(pd, dd):
    # updates points and number of tournaments played
    # pd - player_data, dd - data_dictionary
    dd[pd[1]]['LRT'] += 1
    dd[pd[1]]['PUNKTY'] += tounament_type[pd[0]]

def allocate_points(pd, dd):
    # updates points and number of tournaments played
    # pd - players_data, dd - data_dictionary
    # npn - new_players
    players = [i.split(' ', 1) for i in pd.strip().splitlines()]
    for c, v in enumerate(players):
        if v[1] in dd.keys():
            update_existing_player(v, dd)
        elif v[1] not in dd.keys():
            add_new_player(v, dd)

def ranking(dd):
    # updates ranking place and table traffic
    # pd - player_data, dd - data_dictionary
    ordered_players = []
    for i in dd.keys():
        ordered_players.append([dd[i]['PUNKTY'], dd[i]['LRT'], i])
    ordered_players = sorted(sorted(sorted(ordered_players,
                                key=lambda x: x[2]), 
                                key=lambda x: x[1]), 
                                key=lambda x: x[0], reverse = True)
    l = [ordered_players[0][0], ordered_players[0][1]]
    new_rank = 1
    for c, v in enumerate(ordered_players, 1):
        if v[0] < l[0]:
            l[0], l[1] = v[0], v[1]
            new_rank = c
        elif v[0] == l[0]:
            if v[1] > l[1]:
                l[1] = v[1]
                new_rank = c
        old_rank = int(dd[v[2]]['RANKING'])
        change = old_rank - new_rank
        dd[v[2]]['RANKING'] = new_rank
        if old_rank == 0:
            dd[v[2]]['RUCH'] = '-'
        else:
            if change > 0:
                dd[v[2]]['RUCH'] = '↑ ' + str(abs(change))
            elif change == 0:
                dd[v[2]]['RUCH'] = '-'
            else:
                dd[v[2]]['RUCH'] = '↓ ' + str(abs(change))

# Scoring for cup system
cup = {'1': 40, '2': 30, '4': 20, '8': 15, '16': 10, '32': 8}

# Scoring for 2 groups system
group2 = {'1': 20, '2': 15, '3': 10, '4': 10, '5': 7, '6': 7, '7': 4, 
        '8': 4, '9': 2, '10': 2, '11': 1, '12': 1}

# Scoring for 3 groups system
group3 = {'1': 20, '2': 15, '3': 10, '4': 7, '5': 7, '6': 7, '7': 4, 
        '8': 4, '9': 4, '10': 2, '11': 2, '12': 2}

# Scoring for 4 groups system
group4 = {'1': 20, '2': 15, '3': 10, '4': 10, '5': 7, '6': 7, '7': 7, 
        '8': 7, '9': 4, '10': 4, '11': 4, '12': 4, '13': 2, '14': 2}

# Scoring for Senior Cup
senior = {'1': 20, '2': 15, '3': 10, '4': 7, '5': 4, '6': 0, '7': 0}

# Types of tournament:
# [cup, group2, group3, group4, senior]

tounament_type = cup

# List of tournament participants and their place for eg.
# 1 Andrew Bolton
# 4 Eric Thorn
# etc.

players = """

1 Parszowski Igor
2 Wrona Jakub
4 Dąbrowski Zbigniew
4 Takmadżan Artem
8 Zontek Tomasz
8 Dybka Daniel
8 Maciej Mateusz
8 Wdowiak Tomasz
16 Siek Paweł
16 Pęzioł Mieczysław
16 Kłosowski Szymon
16 Konopelski Paweł
16 Oszga Adam
16 Kadis Jacek
16 Bednarz Damian
16 Marszałek Bartosz
32 Maciej Damian
32 Szeliga Robert
32 Wojtanowski Adrian
32 Michalewski Albert
32 Kosior Kajetan

"""

allocate_points(players, data)
ranking(data)
rows_needed = len(data.keys()) - len(table.rows[1:]) 


def table_creating(dd,tbl, rows_to_add):
    # dd = data_dictionary
    # tbl = table_to_fill
    # rows_needed = new players number,
    #   rows which needs to be added to original table
    # for i in range(rows_to_add):
    #     row = tbl.add_row()
    #     row.height_rule = WD_ROW_HEIGHT.EXACTLY
    #     row.height = Pt(24)
    for i in range(rows_to_add):
        tbl.add_row()
    ordered_players = []
    for i in dd.keys():
        ordered_players.append([dd[i]['PUNKTY'], dd[i]['LRT'], i])
    ordered_players = sorted(sorted(sorted(ordered_players,
                        key=lambda x: x[2]), 
                        key=lambda x: x[1]), 
                        key=lambda x: x[0], reverse = True)
    key_order = ['RANKING', 'RUCH', 'ZAWODNIK', 
                'PUNKTY', 'LRT']
    for i, row in enumerate(tbl.rows[1:]):
        for j, cell in enumerate(row.cells):
            cell.text = str(dd[ordered_players[i][2]][key_order[j]])
    # for i in paragraphs:
    #   print(i.text)
table_creating(data, table, rows_needed)

paragraphs = document.paragraphs
paragraphs[1].text = 'KLASYFIKACJA OPEN MĘŻCZYZN\nNA DZIEŃ ' + datetime.now().strftime('%d.%m.%Y')

document.save(datetime.now().strftime('%d.%m.%Y ') + 'KLASYFIKACJA V GPSWA MĘŻCZYZN.docx')